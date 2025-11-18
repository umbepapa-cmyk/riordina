# Inizio del tuo script Riordina.py

import os
import pdfplumber
import docx
import re
import time
import sys
import openpyxl 
from PIL import Image

# --- NUOVE IMPORTAZIONI PER .ENV e GOOGLE ---
from dotenv import load_dotenv
from google import genai
from google.genai import types
from google.genai.errors import APIError as GeminiAPIError, ServerError as GeminiServerError


# =========================================================================
# 1. CARICAMENTO CONFIGURAZIONE DA .ENV
# =========================================================================
load_dotenv() # Carica le variabili dal file .env

# **⚠️ Le configurazioni ora sono lette dal file .env**
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
CARTELLA_DA_ESAMINARE = os.getenv("CARTELLA_DA_ESAMINARE")

# Valori di default o controlli
GEMINI_MODEL = "gemini-2.5-flash" 
TESTO_MAX_CHAR = 10000   # Limite di caratteri da inviare all'IA
MAX_RETRIES = 5         # Massimo tentativi per la logica di retry manuale

# Controlli di sicurezza per assicurarsi che .env sia stato letto
if not GEMINI_API_KEY:
    print("ERRORE: GEMINI_API_KEY non trovata. Assicurati che sia nel file .env")
    sys.exit(1)
if not CARTELLA_DA_ESAMINARE:
    print("ERRORE: CARTELLA_DA_ESAMINARE non trovata. Assicurati che sia nel file .env")
    sys.exit(1)


# =========================================================================
# 2. FUNZIONE DI ESTRAZIONE DEL TESTO COMPLETO
# =========================================================================

def estrai_testo_con_pypdfium2(percorso_file):
    """
    Funzione di fallback per estrarre testo da PDF usando pypdfium2.
    Utile quando pdfplumber fallisce.
    """
    try:
        import pypdfium2 as pdfium
        
        testo_estratto = ""
        pdf = pdfium.PdfDocument(percorso_file)
        
        for i in range(min(len(pdf), 50)):  # Limite a 50 pagine per performance
            page = pdf[i]
            textpage = page.get_textpage()
            testo_pagina = textpage.get_text_range()
            
            if testo_pagina:
                testo_estratto += testo_pagina + "\n"
                if len(testo_estratto) >= TESTO_MAX_CHAR:
                    break
            
            textpage.close()
            page.close()
        
        pdf.close()
        
        if testo_estratto.strip():
            print(f"   > pypdfium2 ha estratto {len(testo_estratto)} caratteri")
            return testo_estratto
        else:
            print(f"   > pypdfium2 non ha trovato testo (potrebbe essere un PDF basato su immagini)")
            return None
            
    except Exception as e:
        print(f"   > Errore con pypdfium2: {e}")
        return None


def estrai_testo_completo_dal_file(percorso_file):
    """
    Estrae il testo completo (limitato a TESTO_MAX_CHAR) dal file, inclusi i file Excel.
    Usa pdfplumber come metodo primario e pypdfium2 come fallback per i PDF.
    """
    testo_completo = ""
    
    try:
        if percorso_file.lower().endswith(".pdf"):
            # Tentativo primario con pdfplumber
            try:
                with pdfplumber.open(percorso_file) as pdf:
                    if len(pdf.pages) == 0:
                        print(f"   > PDF vuoto o senza pagine: '{os.path.basename(percorso_file)}'")
                        return None
                    
                    for pagina in pdf.pages:
                        testo_pagina = pagina.extract_text()
                        if testo_pagina:
                            testo_completo += testo_pagina + "\n"
                            if len(testo_completo) >= TESTO_MAX_CHAR:
                                break
                    
                    # Se pdfplumber non ha estratto testo, prova con pypdfium2
                    if not testo_completo.strip():
                        print(f"   > pdfplumber non ha estratto testo, provo con pypdfium2...")
                        testo_completo = estrai_testo_con_pypdfium2(percorso_file)
                        
            except Exception as e_pdfplumber:
                print(f"   > pdfplumber fallito: {e_pdfplumber}")
                print(f"   > Tentativo con pypdfium2 come fallback...")
                testo_completo = estrai_testo_con_pypdfium2(percorso_file) 
            
        elif percorso_file.lower().endswith(".docx"):
            try:
                doc = docx.Document(percorso_file)
                for para in doc.paragraphs:
                    testo_completo += para.text + "\n"
                    if len(testo_completo) >= TESTO_MAX_CHAR:
                        break
            except Exception as e:
                print(f"   > Errore lettura DOCX: {e}")
                return None
                    
        elif percorso_file.lower().endswith(".txt"):
            try:
                # Prova diverse codifiche
                for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(percorso_file, 'r', encoding=encoding) as file:
                            testo_completo = file.read(TESTO_MAX_CHAR)
                        break  # Se ha successo, esci dal loop
                    except UnicodeDecodeError:
                        continue
                if not testo_completo:
                    print(f"   > Impossibile decodificare il file TXT con le codifiche comuni")
                    return None
            except Exception as e:
                print(f"   > Errore lettura TXT: {e}")
                return None

        # --- GESTIONE DEI FILE EXCEL (.xlsx) ---
        elif percorso_file.lower().endswith((".xlsx", ".xls")):
            try:
                workbook = openpyxl.load_workbook(percorso_file, read_only=True, data_only=True)
                
                for sheet_name in workbook.sheetnames:
                    foglio = workbook[sheet_name]
                    testo_completo += f"[Foglio: {sheet_name}]\n"
                    
                    for row in foglio.iter_rows():
                        for cell in row:
                            if cell.value is not None:
                                testo_completo += str(cell.value) + " "
                                
                                if len(testo_completo) >= TESTO_MAX_CHAR:
                                    workbook.close()
                                    return testo_completo[:TESTO_MAX_CHAR].strip()
                    testo_completo += "\n"
                
                workbook.close()
            except Exception as e:
                print(f"   > Errore lettura Excel: {e}")
                return None
                        
    except Exception as e:
        print(f"ATTENZIONE: Non ho potuto leggere '{os.path.basename(percorso_file)}'. Motivo: {e}")
        return None
    
    # Gestisci il caso in cui testo_completo sia None (dal fallback pypdfium2)
    if testo_completo is None:
        return None
        
    return testo_completo[:TESTO_MAX_CHAR].strip() if testo_completo.strip() else None


# =========================================================================
# 3. FUNZIONE AVANZATA: Analisi, Titolo (PER TESTO) (Invariata)
# =========================================================================

def analizza_e_titola_gemini(testo_documento, client):
    """
    Invia il testo a Gemini con un prompt migliorato per creare titoli più
    descrittivi ed evitare codici. Gestisce il retry manuale.
    """
    if not testo_documento:
        return None

    # --- PROMPT MIGLIORATO ---
    prompt = f"""
    Sei un assistente di archiviazione file. Il tuo obiettivo è creare un titolo 
    completamente nuovo, pulito e leggibile in italiano, basato sul contenuto di un documento.

    TESTO ESTRATTO DAL DOCUMENTO:
    "{testo_documento}"

    ISTRUZIONI:
    1.  **Analizza il Contenuto:** Leggi il testo e capisci l'argomento principale (es. "scalping", "indici", "strategia V75").
    2.  **Ignora Codici e Nomi File:** Il testo potrebbe contenere vecchi titoli, codici o nomi di file (come 'V75STRATEGYKILLER', '540086610', 'MMTC-Student-s-Manual'). **Ignora queste stringhe letterali.**
    3.  **Crea un Titolo Descrittivo:** Invece di copiare, crea un nuovo titolo che *descriva* l'argomento in modo chiaro e professionale.
        * Se il testo parla di 'V75STRATEGYKILLER', un buon titolo è 'Strategia di Trading V75' o 'Manuale Strategia Indici V75'.
        * Se il testo è 'MMTC-Student-s-Manual-2021', un buon titolo è 'Manuale per Studenti MMTC 2021'.
        * Se il testo è 'Scalping-Presentation', un buon titolo è 'Presentazione sullo Scalping'.
    4.  **Formato:** Il titolo deve essere in italiano, conciso (massimo 8 parole) e pulito.

    Rispondi ESCLUSIVAMENTE con il nuovo titolo descrittivo.
    """
    
    for attempt in range(MAX_RETRIES):
        try:
            response = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=prompt,
                config={"temperature": 0.2} 
            )
            return response.text.strip() 

        except (GeminiServerError, GeminiAPIError) as e:
            if hasattr(e, 'code') and e.code == 503 and attempt < MAX_RETRIES - 1:
                wait_time = 2 ** (attempt + 1)
                print(f"   > [ATTESA 503] Tentativo {attempt + 1}/{MAX_RETRIES} fallito. Attendo {wait_time}s.")
                time.sleep(wait_time)
            else:
                print(f"ERRORE CRITICO (Codice {e.code if hasattr(e, 'code') else 'N/A'}): Richiesta fallita dopo tutti i retry.")
                return None
        except Exception as e:
            print(f"ERRORE di connessione non gestito: {e}")
            return None
    
    return None


# =========================================================================
# 4. FUNZIONE: Analisi, Titolo (PER IMMAGINI) (Corretta)
# =========================================================================

# --- CORREZIONE: Aggiunto 'percorso_file' come argomento ---
def analizza_e_titola_immagine_gemini(percorso_file, client):
    """
    Invia un'immagine a Gemini (multimodale) per ottenere un titolo descrittivo.
    Gestisce il retry manuale.
    """
    try:
        # --- CORREZIONE: Ora 'percorso_file' è definito ---
        img = Image.open(percorso_file)
    except Exception as e:
        print(f"ERRORE: Impossibile aprire l'immagine '{os.path.basename(percorso_file)}'. Motivo: {e}")
        return None

    prompt_immagine = """
    Sei un assistente di archiviazione file. Il tuo obiettivo è creare un titolo 
    completamente nuovo, pulito e leggibile in italiano, basato sul contenuto di questa immagine.

    ISTRUZIONI:
    1.  **Analizza l'Immagine:** Guarda l'immagine e identifica il soggetto, l'azione o il tema principale.
    2.  **Crea un Titolo Descrittivo:** Crea un titolo che *descriva* cosa si vede.
        * Se l'immagine è un grafico azionario, un buon titolo è 'Grafico di Borsa con Medie Mobili'.
        * Se è una foto di un gatto, 'Gatto soriano che dorme'.
        * Se è uno screenshot di un software, 'Screenshot Interfaccia Software'.
    3.  **Formato:** Il titolo deve essere in italiano, conciso (massimo 8 parole) e pulito. Non includere il nome file originale.

    Rispondi ESCLUSIVAMENTE con il nuovo titolo descrittivo.
    """

    for attempt in range(MAX_RETRIES):
        try:
            # La chiamata API multimodale
            response = client.models.generate_content(
                model=GEMINI_MODEL, 
                contents=[prompt_immagine, img], # Lista con prompt e immagine
                config={"temperature": 0.2}
            )
            return response.text.strip() # Successo

        except (GeminiServerError, GeminiAPIError) as e:
            if hasattr(e, 'code') and e.code == 503 and attempt < MAX_RETRIES - 1:
                wait_time = 2 ** (attempt + 1)
                print(f"   > [ATTESA 503] Tentativo {attempt + 1}/{MAX_RETRIES} fallito per immagine. Attendo {wait_time}s.")
                time.sleep(wait_time)
            else:
                print(f"ERRORE API CRITICO (Codice {e.code if hasattr(e, 'code') else 'N/A'}) per immagine. Salto.")
                return None
        except Exception as e:
            print(f"ERRORE di connessione non gestito per immagine: {e}")
            return None
    
    return None # Fallimento dopo tutti i tentativi


# =========================================================================
# 5. INIZIO DELLO SCRIP PRINCIPALE (Inizializzazione da .env)
# =========================================================================

try:
    # L'API Key è già stata caricata e controllata all'inizio
    gemini_client = genai.Client(
        api_key=GEMINI_API_KEY
    )

except Exception as e:
    print("ERRORE FATALE: Impossibile inizializzare il client Gemini. Verifica la tua chiave API nel file .env.")
    print(f"Dettagli: {e}")
    sys.exit(1)


print(f"--- Agente di Rinomina (MODALITÀ AUTOMATICA) avviato ---")
print(f"Cartella sotto osservazione: {CARTELLA_DA_ESAMINARE}")
print(f"NOTA: Analizzo i primi {TESTO_MAX_CHAR} caratteri dei documenti e le immagini.")
print("---")

if not os.path.exists(CARTELLA_DA_ESAMINARE):
    print(f"ERRORE: La cartella specificata '{CARTELLA_DA_ESAMINARE}' non esiste. Controlla il percorso nel file .env.")
    sys.exit(1)

file_rinominati = 0
file_analizzati = 0

# Definisci le estensioni supportate
ESTENSIONI_TESTO = [".pdf", ".docx", ".txt", ".xlsx", ".xls"]
ESTENSIONI_IMMAGINI = [".png", ".jpg", ".jpeg", ".bmp", ".webp"]

# SCANSIONA LA CARTELLA
for nome_file_attuale in os.listdir(CARTELLA_DA_ESAMINARE):
    
    percorso_vecchio = os.path.join(CARTELLA_DA_ESAMINARE, nome_file_attuale)
    
    if not os.path.isfile(percorso_vecchio):
        continue
        
    nome_base, estensione = os.path.splitext(nome_file_attuale)
    estensione = estensione.lower()

    nuovo_nome_base = None

    # --- Blocco 1: Gestione File di Testo ---
    if estensione in ESTENSIONI_TESTO:
        file_analizzati += 1
        testo_documento = estrai_testo_completo_dal_file(percorso_vecchio)
        
        if testo_documento:
            print(f"\n[Analisi Testo] File: '{nome_file_attuale}'")
            print(f"  > Testo inviato: '{testo_documento[:80]}...'")
            nuovo_nome_base = analizza_e_titola_gemini(testo_documento, gemini_client)
            time.sleep(0.5) 
        else:
            print(f"\n[SALTA] Impossibile leggere il testo da '{nome_file_attuale}'")

    # --- Blocco 2: Gestione File Immagine ---
    elif estensione in ESTENSIONI_IMMAGINI:
        file_analizzati += 1
        print(f"\n[Analisi Immagine] File: '{nome_file_attuale}'")
        # --- CORREZIONE: Passiamo 'percorso_vecchio' alla funzione ---
        nuovo_nome_base = analizza_e_titola_immagine_gemini(percorso_vecchio, gemini_client)
        time.sleep(0.5)

    # --- Blocco 3: Logica di Rinomina (Unificata) ---
    if nuovo_nome_base:
        
        # Pulizia e gestione della lunghezza
        nuovo_nome_base_pulito = re.sub(r'[\\/*?:"<>|]', '_', nuovo_nome_base)
        if len(nuovo_nome_base_pulito) > 150:
            nuovo_nome_base_pulito = nuovo_nome_base_pulito[:150].strip()

        # Evita rinomina se il nome è già uguale
        if nuovo_nome_base_pulito.lower() == nome_base.lower():
            print(f"  > Titolo Analizzato: '{nuovo_nome_base_pulito}'")
            print(f"[SALTA] Il nome non richiede modifiche.")
            continue

        # Assemblaggio e gestione duplicati
        nuovo_nome_file = f"{nuovo_nome_base_pulito}{estensione}"
        percorso_nuovo = os.path.join(CARTELLA_DA_ESAMINARE, nuovo_nome_file)
        
        i = 1
        temp_nome_base = nuovo_nome_base_pulito
        while os.path.exists(percorso_nuovo):
            nuovo_nome_file = f"{temp_nome_base}_{i}{estensione}"
            percorso_nuovo = os.path.join(CARTELLA_DA_ESAMINARE, nuovo_nome_file)
            i += 1
        
        try:
            os.rename(percorso_vecchio, percorso_nuovo)
            print(f"  > Titolo Analizzato: '{nuovo_nome_base_pulito}'")
            print(f"[RINOMINATO] '{nome_file_attuale}'  ->  '{nuovo_nome_file}'")
            file_rinominati += 1
        except Exception as e:
            print(f"[ERRORE RINOMINA] Impossibile rinominare '{nome_file_attuale}': {e}")
    
    elif estensione in ESTENSIONI_TESTO or estensione in ESTENSIONI_IMMAGINI:
         print(f"  > L'analisi IA ha restituito un risultato non valido o ha fallito dopo i tentativi. Salto la rinomina.")


print("---")
print(f"Operazione completata.")
print(f"File analizzati (documenti e immagini): {file_analizzati}")
print(f"File rinominati: {file_rinominati}")